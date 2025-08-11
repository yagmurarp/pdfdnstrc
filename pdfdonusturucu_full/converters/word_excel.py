# converters/word_excel.py
from docx import Document
import pandas as pd

def word_to_excel(in_path: str, out_path: str):
    """
    DOCX içindeki tabloları Excel'e yazar.
    Tablonuz yoksa, paragrafları tek sütunda yazar.
    """
    doc = Document(in_path)
    tables = doc.tables
    if tables:
        with pd.ExcelWriter(out_path, engine='openpyxl') as writer:
            for idx, t in enumerate(tables, start=1):
                data = []
                for r in t.rows:
                    data.append([c.text.strip() for c in r.cells])
                df = pd.DataFrame(data)
                df.to_excel(writer, sheet_name=f"Tablo{idx}", index=False, header=False)
    else:
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        df = pd.DataFrame(paras, columns=["Metin"])
        df.to_excel(out_path, index=False)
